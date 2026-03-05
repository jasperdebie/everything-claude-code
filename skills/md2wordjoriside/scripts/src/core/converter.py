import re
from docx import Document
from markdown_it import MarkdownIt
from docx.shared import Pt, RGBColor
from docx.oxml.shared import qn, OxmlElement


class MD2WordConverter:
    def __init__(self, template_path):
        self.template_path = template_path
        self.doc = Document(template_path)
        self.md = MarkdownIt("commonmark", {'breaks': True, 'html': True}).enable('table')

    def parse_markdown(self, md_text):
        return self.md.parse(md_text)

    def _preprocess_markdown(self, text):
        """
        Fixes common issues in Markdown before parsing.
        1. Inserts blank lines before list items that immediately follow text.
        """
        new_text = re.sub(r'\n(\d+\.\s)', r'\n\n\1', text)
        new_text = re.sub(r'(\n+)\d+\.\s+(\*\*.*?\*\*[:]?)\s*(\n)', r'\1\2\3', new_text)
        new_text = re.sub(r'(\n+)[-*+]\s+(\*\*.*?\*\*[:]?)\s*(\n)', r'\n\n\2\3', new_text)
        return new_text

    def generate(self, md_text, output_path, data=None):
        if data:
            self.replace_placeholders(data)
            self.replace_text_in_xml(data)
        
        self.clear_existing_body_content()
        md_text = self._preprocess_markdown(md_text)
        tokens = self.parse_markdown(md_text)
        self.render_tokens(tokens)
        self.set_update_fields()
        self.doc.save(output_path)
        print(f"Generated {output_path}")

    def replace_placeholders(self, data):
        for folder in [self.doc.paragraphs, self.doc.tables]:
            self._replace_in_elements(folder, data)
        for section in self.doc.sections:
            self._replace_in_elements(section.header.paragraphs, data)
            self._replace_in_elements(section.footer.paragraphs, data)

    def _replace_in_elements(self, elements, data):
        if isinstance(elements, list) or hasattr(elements, '__iter__'):
             for elem in elements:
                if hasattr(elem, 'text') and '{' in elem.text:
                    for key, val in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in elem.text:
                            elem.text = elem.text.replace(placeholder, str(val))
                # Also check individual runs (placeholders may be in single runs)
                if hasattr(elem, 'runs'):
                    for run in elem.runs:
                        if run.text and '{' in run.text:
                            for key, val in data.items():
                                placeholder = f"{{{{{key}}}}}"
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(val))
                if hasattr(elem, 'rows'):
                    for row in elem.rows:
                        for cell in row.cells:
                            self._replace_in_elements(cell.paragraphs, data)

    def replace_text_in_xml(self, data):
        xml_sources = [self.doc._element.body]
        for section in self.doc.sections:
             for header in [section.header, section.first_page_header, section.even_page_header]:
                 if header and header._element is not None:
                     xml_sources.append(header._element)
             for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                 if footer and footer._element is not None:
                     xml_sources.append(footer._element)
        
        for xml_element in xml_sources:
             if xml_element is None: continue
             for node in xml_element.iter(qn('w:t')):
                 if not node.text: continue
                 original_text = node.text
                 # Replace any remaining {{placeholder}} patterns in raw XML
                 for key, val in (data or {}).items():
                     placeholder = f"{{{{{key}}}}}"
                     if placeholder in original_text:
                         node.text = node.text.replace(placeholder, str(val))
                         self.unbind_sdt_if_present(node)

    def unbind_sdt_if_present(self, formatting_node):
        current = formatting_node
        sdt_element = None
        for _ in range(6):
            parent = current.getparent()
            if parent is None: break
            if parent.tag == qn('w:sdt'):
                sdt_element = parent
                break
            current = parent
        if sdt_element is not None:
            sdt_content = sdt_element.find(qn('w:sdtContent'))
            if sdt_content is None: return
            parent_of_sdt = sdt_element.getparent()
            index = parent_of_sdt.index(sdt_element)
            for child in list(sdt_content):
                parent_of_sdt.insert(index, child)
                index += 1
            parent_of_sdt.remove(sdt_element)

    def set_update_fields(self):
        body = self.doc._element.body
        fld_chars = list(body.iter(qn('w:fldChar')))
        for fld_char in fld_chars:
            if fld_char.get(qn('w:fldCharType')) in ['begin', 'separate']:
                fld_char.set(qn('w:dirty'), 'true')
        settings = self.doc.settings.element
        update_fields = settings.find(qn('w:updateFields'))
        if update_fields is not None:
             settings.remove(update_fields)

    def clear_existing_body_content(self):
        body = self.doc._element.body
        start_element_index = -1
        for i, child in enumerate(body):
            if child.tag == qn('w:p'):
                pPr = child.find(qn('w:pPr'))
                if pPr is not None:
                    pStyle = pPr.find(qn('w:pStyle'))
                    if pStyle is not None:
                        val = pStyle.get(qn('w:val'))
                        if val and val.startswith('Heading1'):
                            start_element_index = i
                            break
        if start_element_index != -1:
            elements_to_remove = list(body)[start_element_index:]
            for el in elements_to_remove:
                if el.tag == qn('w:sectPr'):
                    continue
                body.remove(el)

    def render_tokens(self, tokens):
        i = 0
        while i < len(tokens):
            token = tokens[i]
            if token.type == 'heading_open':
                level = token.tag
                heading_level = int(level[1])
                style = f'Heading {heading_level}'
                text_token = tokens[i+1]
                if text_token.type == 'inline':
                    text = text_token.content
                    text = re.sub(r'^[\d\.]+\s+', '', text)
                    p = self.doc.add_paragraph(style=style)
                    parts = text.split(r'\n')
                    for part_idx, part in enumerate(parts):
                        if part_idx > 0:
                            p.add_run().add_break()
                        p.add_run(part)
                i += 2
            elif token.type == 'paragraph_open':
                inline_token = tokens[i+1]
                if inline_token.type == 'inline':
                    self.render_inline(inline_token)
                i += 2
            elif token.type == 'bullet_list_open':
                i = self.render_list(tokens, i, list_type='bullet')
            elif token.type == 'ordered_list_open':
                i = self.render_list(tokens, i, list_type='ordered')
            elif token.type == 'table_open':
                i = self.render_table(tokens, i)
            else:
                i += 1

    def render_inline(self, inline_token, paragraph=None):
        if paragraph is None:
            paragraph = self.doc.add_paragraph(style='Normal')
        is_bold = False
        is_italic = False
        for child in inline_token.children:
            if child.type == 'text':
                parts = child.content.split(r'\n')
                for part_idx, part in enumerate(parts):
                    if part_idx > 0:
                        paragraph.add_run().add_break()
                    run = paragraph.add_run(part)
                    if is_bold: run.bold = True
                    if is_italic: run.italic = True
            elif child.type == 'strong_open':
                is_bold = True
            elif child.type == 'strong_close':
                is_bold = False
            elif child.type == 'em_open':
                is_italic = True
            elif child.type == 'em_close':
                is_italic = False
            elif child.type == 'code_inline':
                run = paragraph.add_run(child.content)
                run.font.name = 'Courier New'

    def _get_list_number_abstract_id(self):
        try:
            style = self.doc.styles['List Number']
            pPr = style._element.get_or_add_pPr()
            numPr = pPr.find(qn('w:numPr'))
            numId = numPr.find(qn('w:numId')).get(qn('w:val'))
            numbering_element = self.doc.part.numbering_part.element
            for num in numbering_element.findall(qn('w:num')):
                if num.get(qn('w:numId')) == numId:
                    abstractNumId_elem = num.find(qn('w:abstractNumId'))
                    if abstractNumId_elem is not None:
                        return abstractNumId_elem.get(qn('w:val'))
        except:
            pass
        return None

    def _create_new_list_numbering(self, abstract_num_id, start_val=1):
        numbering_part = self.doc.part.numbering_part
        max_num_id = 0
        for num in numbering_part.element.findall(qn('w:num')):
            nid = int(num.get(qn('w:numId')))
            if nid > max_num_id: max_num_id = nid
        new_num_id = max_num_id + 1
        num_element = OxmlElement('w:num')
        num_element.set(qn('w:numId'), str(new_num_id))
        abstractNumId_element = OxmlElement('w:abstractNumId')
        abstractNumId_element.set(qn('w:val'), str(abstract_num_id))
        num_element.append(abstractNumId_element)
        lvlOverride = OxmlElement('w:lvlOverride')
        lvlOverride.set(qn('w:ilvl'), '0')
        startOverride = OxmlElement('w:startOverride')
        startOverride.set(qn('w:val'), str(start_val))
        lvlOverride.append(startOverride)
        num_element.append(lvlOverride)
        numbering_part.element.append(num_element)
        return new_num_id

    def render_list(self, tokens, i, list_type='bullet'):
        i += 1
        current_num_id = None
        if list_type == 'ordered':
            abstract_id = self._get_list_number_abstract_id()
            if abstract_id:
                current_num_id = self._create_new_list_numbering(abstract_id, start_val=1)

        # Determine best list style available in template
        bullet_style = 'List Bullet'
        number_style = 'List Number'
        # Fallback: check if styles exist, use List Paragraph (present in JI template)
        try:
            self.doc.styles['List Bullet']
        except:
            bullet_style = 'List Paragraph'
        try:
            self.doc.styles['List Number']
        except:
            number_style = 'List Paragraph'

        while i < len(tokens):
            token = tokens[i]
            if token.type == f'{list_type}_list_close':
                return i + 1
            if token.type == 'list_item_open':
                i += 1
                while i < len(tokens) and tokens[i].type != 'list_item_close':
                    if tokens[i].type == 'paragraph_open':
                        inline_t = tokens[i+1]
                        style = bullet_style if list_type == 'bullet' else number_style
                        p = self.doc.add_paragraph(style=style)
                        if list_type == 'ordered' and current_num_id:
                            pPr = p._element.get_or_add_pPr()
                            numPr = pPr.get_or_add_numPr()
                            numId_elem = numPr.find(qn('w:numId'))
                            if numId_elem is None:
                                numId_elem = OxmlElement('w:numId')
                                numPr.append(numId_elem)
                            numId_elem.set(qn('w:val'), str(current_num_id))
                            ilvl_elem = numPr.find(qn('w:ilvl'))
                            if ilvl_elem is None:
                                ilvl_elem = OxmlElement('w:ilvl')
                                numPr.append(ilvl_elem)
                            ilvl_elem.set(qn('w:val'), '0')
                        self.render_inline(inline_t, p)
                        i += 2
                    else:
                        i += 1
                i += 1
            else:
                i += 1
        return i

    def render_table(self, tokens, i):
        rows = []
        current_row = []
        is_header_row = []
        in_thead = False
        i += 1
        while i < len(tokens):
            token = tokens[i]
            if token.type == 'table_close':
                break
            if token.type == 'thead_open':
                in_thead = True
            elif token.type == 'thead_close':
                in_thead = False
            elif token.type == 'tr_open':
                current_row = []
            elif token.type == 'tr_close':
                rows.append(current_row)
                is_header_row.append(in_thead)
            elif token.type in ['th_open', 'td_open']:
                inline = tokens[i+1]
                current_row.append(inline)
            i += 1

        if rows:
            table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))

            # Use Joris Ide template table style hierarchy
            style_applied = False
            for style_name in ['Light List Accent 2', 'Light List - Accent 2', 'Table Grid']:
                try:
                    table.style = style_name
                    style_applied = True
                    break
                except:
                    continue

            for r_idx, row_data in enumerate(rows):
                row = table.rows[r_idx]
                is_header = is_header_row[r_idx] if r_idx < len(is_header_row) else (r_idx == 0)

                for c_idx, inline_token in enumerate(row_data):
                    cell = row.cells[c_idx]
                    p = cell.paragraphs[0]

                    if is_header:
                        # Use tablehead style if available in template
                        try:
                            p.style = self.doc.styles['tablehead']
                        except:
                            pass
                        if inline_token.type == 'inline':
                            self.render_inline(inline_token, p)
                            for run in p.runs:
                                run.font.bold = True
                    else:
                        # Use tabletext style if available in template
                        try:
                            p.style = self.doc.styles['tabletext']
                        except:
                            pass
                        if inline_token.type == 'inline':
                            self.render_inline(inline_token, p)

        return i + 1
